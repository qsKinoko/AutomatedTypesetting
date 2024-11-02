from docx import Document
import os

def modify_bold_text(docx_path, output_path):
    # 打开docx文件
    doc = Document(docx_path)

    # 遍历文档中的所有段落
    for para in doc.paragraphs:
        new_text = ""  # 用来存储整个段落的新文本
        inside_bold = False  # 标记是否在加粗的文本区域

        # 遍历段落中的所有run
        for run in para.runs:
            if run.bold:  # 如果当前run是加粗的
                if not inside_bold:  # 如果还没有进入加粗区域，添加起始的 "|"
                    new_text += "|"
                    inside_bold = True
                new_text += run.text  # 将加粗的文本添加到新文本中
            else:  # 如果当前run不是加粗的
                if inside_bold:  # 如果刚结束加粗区域，添加结束的 "|"
                    new_text += "|"
                    inside_bold = False
                new_text += run.text  # 将非加粗的文本添加到新文本中

        if inside_bold:  # 如果段落结束时仍在加粗区域，关闭 "|"
            new_text += "|"

        # 用新文本替换段落的文本
        para.clear()  # 清空原段落
        para.add_run(new_text)  # 添加新文本

    # 保存修改后的文档
    doc.save(output_path)
    print(f"处理完成！文件已保存为: {output_path}")


# 输入文档路径和输出文档路径
root_path = 'E:\\又填字又填字\\翅刊\\开工！\\v4-110'

input_docx = os.path.join(root_path,'Nightwing v4 #110定稿.docx')
output_docx = os.path.join(root_path,'Nightwing v4 #110定稿_标记.docx')

# 调用函数修改加粗的字
modify_bold_text(input_docx, output_docx)
